import urllib.request
import geocoder
from datetime import date
from bs4 import BeautifulSoup


def get_location():
    g = geocoder.ip('me')
    return g[0]


def get_date():
    today = date.today()
    return today.strftime("%d/%m/%Y")


def get_docx_template():
    from docx import Document
    return Document("/home/aj07mm/template_diario/template_diario.docx")


template_url = "https://docs.google.com/document/d/1xdGwEkjGpa_vpUrHeNTMW1i-Y1B-JGEsILp5csKrhvQ/edit?usp=sharing"
eralegis_url = "https://date.eralegis.info/"

with urllib.request.urlopen(eralegis_url) as response:
   html = response.read()

soup = BeautifulSoup(html, 'html.parser')
eralegis_date = soup.span.contents[0]
cabecalho = """{date} - {location} - {eralegis_date}""".format(
    date=get_date(),
    location=get_location(),
    eralegis_date=eralegis_date,
)

document = get_docx_template()
for paragraph in document.paragraphs:
    if 'cabecalho' in paragraph.text:
        paragraph.clear()
        run = paragraph.add_run()
        run.font.name = "Calibri"
        run.font.underline = True
        run.text = cabecalho
document.save('/home/aj07mm/template_diario/demo.docx')
